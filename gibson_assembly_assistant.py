import questionary
from Bio.SeqUtils import MeltingTemp as mt
from Bio.Seq import Seq
import math
import os
from docx import Document


def annealing_temp_calc():
    fw_primer = None
    rv_primer = None
    knows_tm = questionary.select(
        "Do you know the Tm of your primers?",
        choices=[
            "Yes",
            "No"
        ]
    ).ask()
    if knows_tm == "Yes":
        while True:
            fw_tm = questionary.text("Enter the Tm of your FW primer (in \u00B0C):").ask()
            rv_tm = questionary.text("Enter the Tm of your RV primer (in \u00B0C):").ask()
            try:
                fw_tm = float(fw_tm)
                rv_tm = float(rv_tm)
                break
            except ValueError:
                print("Please enter valid numeric values")
    elif knows_tm == "No":
        while True:
            fw_primer = (questionary.text("Enter the sequence of your FW primer:").ask()).upper()
            rv_primer = (questionary.text("Enter the sequence of your RV primer:").ask()).upper()
            if fw_primer.isalpha() and rv_primer.isalpha():
                fw_tm = mt.Tm_NN(fw_primer, Na=50, dnac1=75)
                rv_tm = mt.Tm_NN(rv_primer, Na=50, dnac1=75)
                break
            else:
                print("Sequences must contain only letters (A, T, C, G)")
    calculated_temp = round((float(fw_tm) + float(rv_tm)) / 2 - 5, 2)
    return(calculated_temp)

def elongation_time_calc(reaction_type):
    insert_vs_plasmid = None
    knows_dna_size = questionary.select(
        "Do you know the size of your DNA sequence?",
        choices=[
            "Yes",
            "No"
        ]
    ).ask()
    if knows_dna_size == "Yes":
        dna_size = int(questionary.text("Enter the DNA size (in bp):").ask())
    elif knows_dna_size == "No":
        script_dir = os.path.dirname(__file__)
        while True:
            user_input = questionary.text(
            "Enter the path to your DNA sequence file (or just the filename):"
        ).ask()
            dna_seq_path = (
                user_input
                if os.path.isabs(user_input)
                else os.path.join(script_dir, user_input)
            )
            if os.path.isfile(dna_seq_path):
                break
            else:
                print("File not found. Please try again.")
        with open(dna_seq_path, 'r') as f:
            dirty_dna_seq = f.read()
        clean_dna_seq = dirty_dna_seq.upper().replace('\r', '').replace('\n', '').replace(' ', '')
        dna_size = int(len(clean_dna_seq))
    elong_time = dna_size/100
    if 60<elong_time<120:
        seconds = elong_time-60
        elong_time = f"1:{round(seconds)}"
    elif elong_time>120:
        seconds = elong_time-120
        elong_time = f"2:{round(seconds)}"
    return elong_time

def dna_prep_amounts(n_samples):
    phanta_ddw_amount = 25*int(n_samples)
    primers_amount = int(n_samples)
    dna_amount = 2*int(n_samples)
    return (phanta_ddw_amount, primers_amount, dna_amount)

def colony_pcr_amounts(n_samples):
    phanta_ddw_amount = 7*int(n_samples)
    primers_amount = int(n_samples)
    return (phanta_ddw_amount, primers_amount)

def dna_size_calculation(seq_path):
    with open(seq_path, 'r') as f:
        dirty_seq = f.read()
    clean_seq = dirty_seq.upper().replace('\r', '').replace('\n', '').replace(' ', '')
    return(clean_seq)

def ligation_amounts():
    knows_vector_size = questionary.select(
        "Do you know the size of your linearized vector?",
        choices=[
            "Yes",
            "No"
        ]
    ).ask()
    if knows_vector_size == "Yes":
        lin_vector_size = int(questionary.text("Enter the linearized vector size (in bp):").ask())
    elif knows_vector_size == "No":
        script_dir = os.path.dirname(__file__)
        while True:
            user_input = questionary.text(
            "Enter the path to a text file containing the linearized vector sequence:").ask()
            lin_vector_seq_path = (
                user_input
                if os.path.isabs(user_input)
                else os.path.join(script_dir, user_input)
            )
            if os.path.isfile(lin_vector_seq_path):
                break
            else:
                print("File not found. Please try again.")
        lin_vector_size = int(len(dna_size_calculation(lin_vector_seq_path)))
    vector_conc = float(questionary.text("Enter the concentration of the linearized vector (in ng/µl):").ask())
    knows_insert_size = questionary.select(
        "Do you know the size of your insert (with overhangs)?",
        choices=[
            "Yes",
            "No"
        ]
    ).ask()
    if knows_insert_size == "Yes":
        insert_size = int(questionary.text("Enter the insert size (in bp):").ask())
    elif knows_insert_size == "No":
        script_dir = os.path.dirname(__file__)
        while True:
            user_input = questionary.text(
            "Enter the path to a text file containing the insert with overhangs sequence:").ask()
            insert_seq_path = (
                user_input
                if os.path.isabs(user_input)
                else os.path.join(script_dir, user_input)
            )
            if os.path.isfile(insert_seq_path):
                break
            else:
                print("File not found. Please try again.")
        insert_size = int(len(dna_size_calculation(insert_seq_path)))
    insert_conc = float(questionary.text("Enter the concentration of the insert (in ng/µl):").ask())
    needed_vector_vol = (6*(lin_vector_size/1000))/vector_conc
    needed_insert_vol = (12*(insert_size/1000))/insert_conc

    dilution_info = []

    def normalize(vol):
        if vol < 1:
            dilution_factor = math.ceil(1 / vol)
            return (vol*dilution_factor), dilution_factor
        return vol, 1

    plasmid_final, plasmid_dilution = normalize(needed_vector_vol)
    insert_final, insert_dilution = normalize(needed_insert_vol)
    total = plasmid_final + insert_final
    reduction_note = None

    if total > 3:
        excess = total - 3
        if plasmid_final >= insert_final:
            reduced = plasmid_final - excess
            reduction_note = f"Vector concentration reduced from {plasmid_final:.2f} to {reduced:.2f} ng/µL to keep total at 3 ng/µL"
            plasmid_final = reduced
        else:
            reduced = insert_final - excess
            reduction_note = f"Insert concentration reduced from {insert_final:.2f} to {reduced:.2f} ng/µL to keep total at 3 ng/µL"
            insert_final = reduced
    
    final_vec_vol = round(plasmid_final, 2)
    final_ins_vol = round(insert_final, 2)

    return (final_vec_vol, final_ins_vol, plasmid_dilution, insert_dilution, reduction_note)

def fill_docx_template(docx_path, output_path, replacements):
    doc = Document(docx_path)

    def replace_placeholders_in_paragraph(paragraph):
        # Join all run texts to one string
        full_text = ''.join(run.text for run in paragraph.runs)

        for key, val in replacements.items():
            full_text = full_text.replace(f"{{{key}}}", str(val))

        if paragraph.runs:
            paragraph.clear()
        paragraph.add_run(full_text)

    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph)

    doc.save(output_path)

def protocol_generator(reaction_type, experiment_name, summary, n_samp=None, lin_or_overhangs=None):
    script_dir = os.path.dirname(__file__)
    template_paths = {
        "Plasmid linearization / Insert preparation": os.path.join(script_dir, "lin_ins_protocol.docx"),
        "Colony PCR": os.path.join(script_dir, "colony_pcr_protocol.docx"),
        "Ligation": os.path.join(script_dir, "ligation_protocol.docx")
    }
    template_path = template_paths[reaction_type]
    want_protocol = questionary.select(
        "Do you want to save a printable protocol?",
        choices=[
            "Yes",
            "No"
        ]
    ).ask()

    if want_protocol == "No":
        print("No protocol generated")
        return None

    if want_protocol == "Yes":
        if reaction_type == "Plasmid linearization / Insert preparation":
            if lin_or_overhangs == None:
                lin_or_overhangs = questionary.select(
                "Are you linearizing a plasmid or preparing an insert with overhangs?",
                choices=[
                    "Linearizing a plasmid",
                    "Preparing an insert"
                ]
            ).ask()
            replacements = {
            "lin_or_oh": lin_or_overhangs,
            "exp_name": experiment_name,
            "n": n_samp,
            "Phanta": summary["Phanta master mix"],
            "DDW": summary["ddH2O"],
            "FW": summary["Forward primer (undiluted)"],
            "RV": summary["Reverse primer (undiluted)"],
            "DNA": summary["DNA amount"],
            "anneal": summary["Annealing temperature"],
            "elongate": summary["Elongation time"]
        }
            
        elif reaction_type == "Colony PCR":
            replacements = {
            "exp_name": experiment_name,
            "n": n_samp,
            "Phanta": summary["Phanta master mix"],
            "DDW": summary["ddH2O"],
            "FW": summary["Forward primer (1:10 dilution)"],
            "RV": summary["Reverse primer (1:10 dilution)"],
            "anneal": summary["Annealing temperature"],
            "elongate": summary["Elongation time"]
        }
            
        elif reaction_type == "Ligation":
            replacements = {
            "exp_name": experiment_name,
            "vec_dil": summary["Vector dilution"],
            "vec": summary["Vector"],
            "ins_dil": summary["Insert dilution"],
            "ins": summary["Insert"],
            "reduction_note": summary["Attention"]
        }
      
    output_path = os.path.join(script_dir, f"{experiment_name}_protocol.docx")
    fill_docx_template(template_path, output_path, replacements)
    return output_path   

#-------------------------------------------------------------------------------------------------------------

def main():
    experiment_name = questionary.text("Please type the name of your experiment:").ask()
    reaction_type = questionary.select(
        "What type of PCR reaction are you planning?",
        choices=[
            "Plasmid linearization / Insert preparation",
            "Ligation",
            "Colony PCR"
        ]
    ).ask()

    if reaction_type == "Plasmid linearization / Insert preparation":
        annealing_temp = annealing_temp_calc()
        elongation_time = elongation_time_calc(reaction_type)
        n_samp = int(questionary.text("Enter the number of samples:").ask())
        phanta_ddw_amount, primers_amount, dna_amount = dna_prep_amounts(n_samp)
        if isinstance(elongation_time, (int, float)):
            units = "sec"
        else:
            units = "min"
        insert_vs_plasmid = questionary.select(
                "Are you linearizing a plasmid or preparing an insert with overhangs?",
                choices=[
                    "Linearizing a plasmid",
                    "Preparing an insert"
                ]
            ).ask()
        summary = {
            "Annealing temperature" : f"{annealing_temp}°C",
            "Elongation time" : f"{elongation_time} {units}",
            "Phanta master mix" : f"{phanta_ddw_amount} µL",
            "ddH2O" : f"{phanta_ddw_amount} µL",
            "Forward primer (undiluted)" : f"{primers_amount} µL",
            "Reverse primer (undiluted)" : f"{primers_amount} µL",
            "DNA amount" : f"{dna_amount} µL"
        }
        for key, value in summary.items():
            print(f"{key:<35} {value}")
        generated_protocol = protocol_generator(reaction_type, experiment_name, summary, n_samp, insert_vs_plasmid)
        if generated_protocol != None:
            print(f"Protocol saved as {generated_protocol}")
        start_again = questionary.select(
            "Do you need another calculation?",
            choices=[
                "Yes",
                "No"
            ]
        ).ask()
        if start_again == "Yes":
            main()
        else:
            print("Exits program")
            return None


    elif reaction_type == "Colony PCR":
        annealing_temp = annealing_temp_calc()
        elongation_time = elongation_time_calc(reaction_type)
        n_samp = int(questionary.text("Enter the number of samples:").ask())
        phanta_ddw_amount, primers_amount = colony_pcr_amounts(n_samp)
        if isinstance(elongation_time, (int, float)):
            units = "sec"
        else:
            units = "min"
        summary = {
            "Annealing temperature" : f"{annealing_temp}°C",
            "Elongation time" : f"{elongation_time} {units}",
            "Phanta master mix" : f"{phanta_ddw_amount} µL",
            "ddH2O" : f"{phanta_ddw_amount} µL",
            "Forward primer (1:10 dilution)" : f"{primers_amount} µL",
            "Reverse primer (1:10 dilution)" : f"{primers_amount} µL",
            "Colony" : "Add manually"
        }
        for key, value in summary.items():
            print(f"{key:<35} {value}")
        generated_protocol = protocol_generator(reaction_type, experiment_name, summary, n_samp)
        if generated_protocol != None:
            print(f"Protocol saved as {generated_protocol}")
        start_again = questionary.select(
            "Do you need another calculation?",
            choices=[
                "Yes",
                "No"
            ]
        ).ask()
        if start_again == "Yes":
            main()
        else:
            print("Exits program")
            return None

    elif reaction_type == "Ligation":
        final_vec_conc, final_ins_conc, plasmid_dilution, insert_dilution, reduction_note = ligation_amounts()
        summary = {
            "Enzyme mix" : "7.5 µL",
            "Vector dilution" : f"1:{plasmid_dilution}",
            "Insert dilution" : f"1:{insert_dilution}",
            "Vector" : f"{final_vec_conc} µL",
            "Insert" : f"{final_ins_conc} µL",
            "Attention" : reduction_note
        }
        for key, value in summary.items():
            print(f"{key:<35} {value}")
        generated_protocol = protocol_generator(reaction_type, experiment_name, summary)
        if generated_protocol != None:
            print(f"Protocol saved as {generated_protocol}")
        start_again = questionary.select(
            "Do you need another calculation?",
            choices=[
                "Yes",
                "No"
            ]
        ).ask()
        if start_again == "Yes":
            main()
        else:
            print("Exits program")
            return None

#Beginning
if __name__ == "__main__":
    main()
